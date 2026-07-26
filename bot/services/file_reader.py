from __future__ import annotations

import io
import logging
import re
from base64 import b64encode

import pdfplumber
import pypdfium2 as pdfium
from docx import Document as DocxDocument
from PIL import Image

logger = logging.getLogger(__name__)

CHUNK_SIZE = 4000
SCAN_THRESHOLD = 100


def extract_text(file_bytes: bytes, ext: str) -> str:
    ext = ext.lower()
    if ext == "pdf":
        return _extract_pdf(file_bytes)
    if ext in ("docx",):
        return _extract_docx(file_bytes)
    raise ValueError(f"Unsupported format: {ext}")


def is_scanned_pdf(file_bytes: bytes) -> bool:
    text = _extract_pdf(file_bytes)
    clean = re.sub(r"--- Страница \d+ ---\s*", "", text).strip()
    return len(clean) < SCAN_THRESHOLD


def pdf_pages_as_base64(file_bytes: bytes, max_pages: int = 10) -> list[str]:
    pdf = pdfium.PdfDocument(io.BytesIO(file_bytes))
    total = min(len(pdf), max_pages)
    images: list[str] = []

    for i in range(total):
        page = pdf[i]
        bitmap = page.render(scale=2)
        pil = Image.frombytes("RGB", (bitmap.width, bitmap.height), bitmap.format("bgr"))
        buf = io.BytesIO()
        pil.save(buf, format="JPEG", quality=85)
        images.append(b64encode(buf.getvalue()).decode())

    pdf.close()
    return images


def _extract_pdf(file_bytes: bytes) -> str:
    pages = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            pages.append(f"--- Страница {i} ---\n{text.strip()}")
    return "\n\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    return "\n".join(paragraphs)


def chunk_text(text: str, size: int = CHUNK_SIZE) -> list[str]:
    paragraphs = text.split("\n\n")
    chunks: list[str] = []
    current = ""

    for para in paragraphs:
        if not para.strip():
            continue
        if len(current) + len(para) < size:
            current += "\n\n" + para if current else para
        else:
            if current:
                chunks.append(current.strip())
            current = para

    if current:
        chunks.append(current.strip())

    return chunks or ["(пустой файл)"]
