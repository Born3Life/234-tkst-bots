from __future__ import annotations

import io
import re

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_answer_docx(text: str) -> io.BytesIO:
    doc = DocxDocument()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    DARK_BLUE = RGBColor(0x00, 0x2B, 0x5C)

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            i += 1
            continue

        heading = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if heading:
            level = len(heading.group(1))
            title = heading.group(2)
            h = doc.add_heading(title, level=min(level, 3))
            for run in h.runs:
                run.font.color.rgb = DARK_BLUE
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, stripped[2:], DARK_BLUE)
            i += 1
            continue

        num = re.match(r'^(\d+)[.)]\s+(.+)$', stripped)
        if num:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, num.group(2), DARK_BLUE)
            i += 1
            continue

        p = doc.add_paragraph()
        _add_runs(p, lines[i], DARK_BLUE)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_runs(paragraph, text: str, dark_blue: RGBColor) -> None:
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = dark_blue
        elif part.startswith("*") and part.endswith("*") and len(part) > 2 and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)
