from __future__ import annotations

import io
import re

from docx import Document
from docx.shared import Pt


def create_answer_docx(text: str) -> io.BytesIO:
    doc = Document()

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue

        if line.startswith("# ") or (line.startswith("**") and line.endswith("**")):
            heading = line.strip("*# ")
            doc.add_heading(heading, level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- ") or line.startswith("• "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+[.)]", line):
            doc.add_paragraph(line, style="List Number")
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
